#!/usr/bin/env -S uv run --script
# /// script
# requires-python = ">=3.11"
# dependencies = [
#   "python-docx>=1.1.0",
#   "markitdown>=0.0.1a2",
#   "docx2txt>=0.8",
# ]
# ///
"""
Parse Word documents using 4 extraction methods: basic text+images, fast text-only,
detailed with tables/headings, and simple text extraction.

Usage: ./parse_docx.py <file_path> <output_dir>
"""

import sys
import json
from pathlib import Path
from datetime import datetime
from docx import Document
from markitdown import MarkItDown
import docx2txt


def parse_method1_python_docx(source_file, output_dir):
    """Extract text and images using python-docx."""
    print("    Method 1: python-docx (basic)...")
    method_dir = output_dir / "python_docx_basic"
    method_dir.mkdir(exist_ok=True)
    images_dir = method_dir / "images"
    images_dir.mkdir(exist_ok=True)

    try:
        doc = Document(source_file)

        text_content = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_content.append(para.text)

        with open(method_dir / "content.md", "w") as f:
            f.write("\n\n".join(text_content))

        img_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                img_count += 1
                img_data = rel.target_part.blob
                ext = rel.target_ref.split(".")[-1]
                with open(images_dir / f"image_{img_count:03d}.{ext}", "wb") as f:
                    f.write(img_data)

        print(
            f"      ✓ Extracted {len(doc.paragraphs)} paragraphs and {img_count} images"
        )
        return True
    except Exception as e:
        print(f"      ⚠ Failed: {str(e)[:100]}")
        return False


def parse_method2_markitdown(source_file, output_dir):
    """Fast text-only extraction using Microsoft's markitdown."""
    print("    Method 2: markitdown...")
    method_dir = output_dir / "markitdown"
    method_dir.mkdir(exist_ok=True)

    try:
        md = MarkItDown()
        result = md.convert(str(source_file))

        with open(method_dir / "content.md", "w") as f:
            f.write(result.text_content)

        print(f"      ✓ Text extracted")
        return True
    except Exception as e:
        print(f"      ⚠ Failed: {str(e)[:100]}")
        return False


def parse_method3_python_docx_detailed(source_file, output_dir):
    """Extract with heading styles, tables as JSON, and document metadata."""
    print("    Method 3: python-docx (detailed)...")
    method_dir = output_dir / "python_docx_detailed"
    method_dir.mkdir(exist_ok=True)

    try:
        doc = Document(source_file)

        metadata = {
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "sections": len(doc.sections),
        }

        if hasattr(doc, "core_properties"):
            core_props = doc.core_properties
            core_props_data = {}
            if hasattr(core_props, "author"):
                core_props_data["author"] = core_props.author
            if hasattr(core_props, "title"):
                core_props_data["title"] = core_props.title
            if hasattr(core_props, "created"):
                core_props_data["created"] = str(core_props.created)
            if hasattr(core_props, "modified"):
                core_props_data["modified"] = str(core_props.modified)
            metadata["core_properties"] = core_props_data

        with open(method_dir / "metadata.json", "w") as f:
            json.dump(metadata, f, indent=2)

        content = []

        for para in doc.paragraphs:
            if para.text.strip():
                style = para.style.name if para.style else "Normal"
                if "Heading" in style:
                    level = style.replace("Heading ", "").strip()
                    try:
                        heading_level = int(level)
                        content.append(f"{'#' * heading_level} {para.text}")
                    except:
                        content.append(f"## {para.text}")
                else:
                    content.append(para.text)

        for table_idx, table in enumerate(doc.tables, 1):
            content.append(f"\n### Table {table_idx}\n")
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)

            with open(method_dir / f"table_{table_idx}.json", "w") as f:
                json.dump(table_data, f, indent=2)

            if table_data:
                content.append("| " + " | ".join(table_data[0]) + " |")
                content.append("|" + "---|" * len(table_data[0]))
                for row in table_data[1:]:
                    content.append("| " + " | ".join(row) + " |")

        with open(method_dir / "content.md", "w") as f:
            f.write("\n\n".join(content))

        print(
            f"      ✓ Extracted {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables"
        )
        return True
    except Exception as e:
        print(f"      ⚠ Failed: {str(e)[:100]}")
        return False


def parse_method4_docx2txt(source_file, output_dir):
    """Lightweight text extraction with images using docx2txt."""
    print("    Method 4: docx2txt...")
    method_dir = output_dir / "docx2txt"
    method_dir.mkdir(exist_ok=True)
    images_dir = method_dir / "images"
    images_dir.mkdir(exist_ok=True)

    try:
        text = docx2txt.process(str(source_file), str(images_dir))

        with open(method_dir / "content.txt", "w") as f:
            f.write(text)

        img_count = len(list(images_dir.glob("*")))

        print(f"      ✓ Extracted text and {img_count} images")
        return True
    except Exception as e:
        print(f"      ⚠ Failed: {str(e)[:100]}")
        return False


def main():
    if len(sys.argv) != 3:
        print("Usage: ./parse_docx.py <file_path> <output_dir>")
        sys.exit(1)

    source_file = Path(sys.argv[1])
    output_dir = Path(sys.argv[2])

    if not source_file.exists():
        print(f"Error: File not found: {source_file}")
        sys.exit(1)

    print(f"Parsing DOCX: {source_file.name}")
    print(f"Output directory: {output_dir}")
    print()

    output_dir.mkdir(parents=True, exist_ok=True)

    methods_success = {
        "method1_python_docx_basic": parse_method1_python_docx(source_file, output_dir),
        "method2_markitdown": parse_method2_markitdown(source_file, output_dir),
        "method3_python_docx_detailed": parse_method3_python_docx_detailed(
            source_file, output_dir
        ),
        "method4_docx2txt": parse_method4_docx2txt(source_file, output_dir),
    }

    success_count = sum(methods_success.values())
    print(f"\n  ✓ Completed: {success_count}/{len(methods_success)} methods successful")

    results = {
        "timestamp": datetime.now().isoformat(),
        "source_file": str(source_file),
        "output_dir": str(output_dir),
        "methods": methods_success,
        "success_count": success_count,
    }

    with open(output_dir / "parsing_results.json", "w") as f:
        json.dump(results, f, indent=2)

    return 0 if success_count > 0 else 1


if __name__ == "__main__":
    sys.exit(main())
